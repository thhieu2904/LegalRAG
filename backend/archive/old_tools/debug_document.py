#!/usr/bin/env python3
"""
Debug Document Text Structure
"""

import sys
from pathlib import Path
from docx import Document

# Add backend to path
sys.path.append(str(Path(__file__).parent))

def debug_document_structure():
    print("🔍 DEBUGGING DOCUMENT STRUCTURE")
    print("=" * 50)
    
    # Find sample document
    documents_dir = Path("data/documents")
    sample_file = None
    
    for collection_dir in documents_dir.glob("*"):
        if collection_dir.is_dir():
            for file in collection_dir.glob("*.doc*"):
                if file.is_file():
                    sample_file = str(file)
                    break
            if sample_file:
                break
    
    if not sample_file:
        print("❌ No sample document found!")
        return
    
    print(f"📄 Analyzing: {Path(sample_file).name}")
    print("-" * 40)
    
    try:
        doc = Document(sample_file)
        
        print(f"Total paragraphs: {len(doc.paragraphs)}")
        
        # Analyze first 10 paragraphs
        for i, para in enumerate(doc.paragraphs[:10]):
            text = para.text.strip()
            if text:
                print(f"\nParagraph {i+1}: {len(text)} chars")
                print(f"  Content: {repr(text[:100])}{'...' if len(text) > 100 else ''}")
                print(f"  Ends with: {repr(text[-10:]) if len(text) > 10 else repr(text)}")
        
        # Check how paragraphs are being joined
        print(f"\n📋 TEXT EXTRACTION METHODS:")
        print("-" * 30)
        
        # Method 1: Direct join
        method1 = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                method1.append(text)
        
        result1 = '\n'.join(method1)
        
        print(f"Method 1 (\\n join): {len(result1)} chars, {result1.count(chr(10))} newlines")
        print(f"  First 200 chars: {repr(result1[:200])}")
        
        # Method 2: Smart joining with paragraph detection
        method2_parts = []
        current_section = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            # Check if this looks like a section header
            if (text.isupper() and len(text) < 100) or text.endswith(':') or any(text.startswith(x) for x in ['1.', '2.', '3.', 'a)', 'b)', 'c)']):
                # Save previous section
                if current_section:
                    method2_parts.append('\n'.join(current_section))
                # Start new section
                current_section = [text]
            else:
                current_section.append(text)
        
        # Add final section
        if current_section:
            method2_parts.append('\n'.join(current_section))
        
        result2 = '\n\n'.join(method2_parts)
        
        print(f"Method 2 (section-aware): {len(result2)} chars, {result2.count(chr(10))} newlines")
        print(f"  First 200 chars: {repr(result2[:200])}")
        print(f"  Sections found: {len(method2_parts)}")
        
        # Show section structure
        print(f"\n📋 SECTION STRUCTURE:")
        for i, section in enumerate(method2_parts[:5]):
            lines = section.split('\n')
            print(f"  Section {i+1}: {len(lines)} lines, {len(section)} chars")
            if lines:
                print(f"    Header: {repr(lines[0][:50])}...")
    
    except Exception as e:
        print(f"❌ Error analyzing document: {e}")

if __name__ == "__main__":
    debug_document_structure()
